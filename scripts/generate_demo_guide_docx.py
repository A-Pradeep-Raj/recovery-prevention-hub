"""Generates a Word (.docx) version of demo_guide.md with embedded screenshots.

This script is a local-only convenience tool — the generated .docx is NOT
committed to git (see .gitignore: Demo_Guide_*.docx). Run it any time the
markdown demo guide or screenshots change:

    cd backend
    .\\venv\\Scripts\\Activate.ps1
    python ..\\scripts\\generate_demo_guide_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "demo_guide.md"
OUT_PATH = ROOT / "Demo_Guide_Recovery_Prevention_Hub.docx"
SCREENSHOTS_DIR = ROOT / "docs" / "screenshots"

HEADING_COLOR = RGBColor(0x1F, 0x4E, 0x79)


def add_title_page(doc: Document) -> None:
    title = doc.add_heading("Recovery & Prevention Hub", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Live Demo Guide — Judged Session Walkthrough")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.color.rgb = HEADING_COLOR
    meta = doc.add_paragraph("Generated from demo_guide.md — includes live screenshots captured 2026-07-25")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].italic = True
    url = doc.add_paragraph("https://recovery-hub-frontend-566288522012.us-central1.run.app/")
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_markdown_line(doc: Document, line: str) -> None:
    line = line.rstrip("\n")

    # Image syntax: ![alt](path)
    img_match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if img_match:
        alt_text, rel_path = img_match.groups()
        img_path = ROOT / rel_path
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption = doc.add_paragraph(alt_text)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].italic = True
            caption.runs[0].font.size = Pt(9)
        else:
            doc.add_paragraph(f"[Missing image: {rel_path}]")
        return

    # Headings
    heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
    if heading_match:
        hashes, text = heading_match.groups()
        level = len(hashes)
        text = re.sub(r"[*`]", "", text)
        doc.add_heading(text, level=min(level, 4))
        return

    # Blockquote
    if line.strip().startswith(">"):
        text = line.strip().lstrip(">").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        return

    # Horizontal rule / empty line
    if line.strip() in ("", "---"):
        return

    # Table row (very simple pipe-table handling is done separately below)
    if line.strip().startswith("|"):
        return  # handled by table block processor

    # List item
    list_match = re.match(r"^\s*[-*]\s+(.*)", line)
    if list_match:
        text = list_match.group(1)
        p = doc.add_paragraph(style="List Bullet")
        add_inline_runs(p, text)
        return

    checklist_match = re.match(r"^\s*-\s+\[( |x|X)\]\s+(.*)", line)
    if checklist_match:
        checked, text = checklist_match.groups()
        box = "[x]" if checked.lower() == "x" else "[ ]"
        p = doc.add_paragraph(style="List Bullet")
        add_inline_runs(p, f"{box} {text}")
        return

    # Regular paragraph
    p = doc.add_paragraph()
    add_inline_runs(p, line)


def add_inline_runs(paragraph, text: str) -> None:
    """Very small inline markdown renderer: **bold**, *italic*, `code`."""
    token_re = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    pos = 0
    for m in token_re.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB3, 0x00, 0x00)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    # Drop the markdown separator row (e.g. |---|---|)
    if len(rows) >= 2 and all(re.match(r"^:?-+:?$", c) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Light Grid Accent 1"
    for row_data in rows:
        row_cells = table.add_row().cells
        for idx, cell_text in enumerate(row_data[:n_cols]):
            row_cells[idx].text = re.sub(r"[*`]", "", cell_text)


def build_document() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # Base styling
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_page(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|"):
            rows, next_i = parse_table_block(lines, i)
            add_table(doc, rows)
            i = next_i
            continue
        if line.strip().startswith("```"):
            # Skip fenced code blocks' delimiter, render contents as monospace paragraph
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            continue
        add_markdown_line(doc, line)
        i += 1

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build_document()
