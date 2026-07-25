"""Generates a SHORT, presenter-friendly Word (.docx) demo script for judges —
big steps, one line of instruction each, one screenshot each. This is a local-
only file (see .gitignore: *.docx) and is never pushed to GitHub.

Run:
    cd backend
    .\\venv\\Scripts\\Activate.ps1
    python ..\\scripts\\generate_simple_demo_docx.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
SCREENSHOTS_DIR = ROOT / "docs" / "screenshots"
OUT_PATH = ROOT / "Live_Demo_For_Judges.docx"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
GRAY = RGBColor(0x55, 0x55, 0x55)

STEPS = [
    {
        "num": 1,
        "title": "Open the App",
        "say": "This is the Recovery & Prevention Hub, live on Google Cloud.",
        "do": "Open the URL below. You land directly on Crisis Mode — no login screen, no menus.",
        "image": "01-crisis-mode-landing.png",
    },
    {
        "num": 2,
        "title": "Build a Real Safety Plan",
        "say": "Everything the AI generates later comes only from what we type here — nothing invented.",
        "do": "Go to the Safety Plan tab. Type a trigger, a coping strategy, and a support contact. Click Save.",
        "image": "03-safety-plan.png",
    },
    {
        "num": 3,
        "title": "Tap the Crisis Button — Zero Typing",
        "say": "One tap. Gemini generates this live, right now, grounded in what we just typed.",
        "do": "Go to Crisis Mode. Tap the big \"I need help now\" button. Wait a few seconds.",
        "image": "02-crisis-mode-grounded-result.png",
    },
    {
        "num": 4,
        "title": "Craving Check-In",
        "say": "No typing here either — just tap an intensity level.",
        "do": "In the Safety Plan tab, scroll down and tap any intensity button (e.g. 8/10).",
        "image": "04-craving-checkin.png",
    },
    {
        "num": 5,
        "title": "Ask the Recovery Co-Pilot — Grounded Question",
        "say": "The answer comes only from the Safety Plan we saved — with a source shown.",
        "do": "Go to Recovery Co-Pilot. Type: \"What coping strategies have worked for me before?\" Click Ask.",
        "image": "05-copilot-grounded.png",
    },
    {
        "num": 6,
        "title": "Ask an Unsafe Question — Watch It Refuse",
        "say": "This is the safety guardrail: it never invents a medical answer.",
        "do": "In the same tab, type: \"What medication dosage should I take for withdrawal?\" Click Ask.",
        "image": "06-copilot-guardrail.png",
    },
    {
        "num": 7,
        "title": "Caregiver Dashboard",
        "say": "The caregiver sees this in real time — with a suggested next action.",
        "do": "Go to Caregiver Dashboard. Show the new alert. Click Acknowledge.",
        "image": "07-caregiver-dashboard.png",
    },
]


def add_cover(doc: Document) -> None:
    title = doc.add_heading("Recovery & Prevention Hub", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("LIVE DEMO — For Judges")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(20)
    sub.runs[0].font.bold = True
    sub.runs[0].font.color.rgb = BLUE

    sub2 = doc.add_paragraph("7 simple steps. Every result shown is generated live by real Gemini AI.")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].italic = True

    url = doc.add_paragraph()
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = url.add_run("https://recovery-hub-frontend-566288522012.us-central1.run.app/")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = GREEN

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Tip: Follow the steps in order. Each step shows what to click, what to say, "
        "and a screenshot of the expected result."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(10)
    note.runs[0].font.color.rgb = GRAY

    doc.add_page_break()


def add_step(doc: Document, step: dict) -> None:
    heading = doc.add_heading(f"Step {step['num']} — {step['title']}", level=1)
    heading.runs[0].font.color.rgb = BLUE

    do_p = doc.add_paragraph()
    do_p.add_run("DO: ").bold = True
    do_p.add_run(step["do"])
    do_p.paragraph_format.space_after = Pt(4)

    say_p = doc.add_paragraph()
    say_run = say_p.add_run("SAY: ")
    say_run.bold = True
    say_run.font.color.rgb = GREEN
    quote_run = say_p.add_run(f"\u201c{step['say']}\u201d")
    quote_run.italic = True
    say_p.paragraph_format.space_after = Pt(10)

    img_path = SCREENSHOTS_DIR / step["image"]
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Missing screenshot: {step['image']}]")

    if step["num"] < len(STEPS):
        doc.add_page_break()


def add_closing(doc: Document) -> None:
    doc.add_heading("Closing Line", level=1).runs[0].font.color.rgb = BLUE
    p = doc.add_paragraph()
    run = p.add_run(
        "\u201cEverything you just saw — the crisis script, the check-in suggestion, the "
        "chatbot answers — was generated live, right now, by real Gemini AI. Nothing was "
        "pre-written. The safety guardrails mean it either uses your real data, or it "
        "honestly says it doesn't know.\u201d"
    )
    run.italic = True
    run.font.size = Pt(13)


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    add_cover(doc)
    for step in STEPS:
        add_step(doc, step)
    doc.add_page_break()
    add_closing(doc)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
